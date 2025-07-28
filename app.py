import streamlit as st
import pandas as pd
import io
import os
from openai import OpenAI
from openai import AuthenticationError, APIConnectionError, RateLimitError, APIStatusError # Import specific OpenAI errors
import requests # Import requests for potential timeout errors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn # For setting table style
from docx.oxml import OxmlElement # For setting table style
import matplotlib.pyplot as plt
import seaborn as sns
import re # For parsing graph requests and markdown bolding
from scipy import stats # For statistical tests
import numpy as np # For numerical operations, especially for ANOVA SS calculations
import pingouin as pg # NEW: For Cronbach's Alpha
import statsmodels.api as sm # NEW: For Linear Regression
import statsmodels.formula.api as smf # NEW: For Linear Regression with formula API

# --- Configuration and Secrets ---
# IMPORTANT: Create a .streamlit/secrets.toml file in your project root
# and add your OpenAI API key and login credentials there.
# Example secrets.toml structure:
# [openai]
# api_key = "sk-proj-your_openai_api_key_here"
#
# [credentials]
# user1 = "User1@123"
# user2 = "User2@123"
# "ssoconsultants14@gmail.com" = "Sso@122" # Corrected password for ssoconsultants14@gmail.com


# Load secrets
try:
    OPENAI_API_KEY = st.secrets["openai"]["api_key"] # Accessing the api_key from the [openai] section
    LOGIN_USERS = st.secrets["credentials"]          # Accessing the users from the [credentials] section
except KeyError as e:
    st.error(f"Secret not found: {e}. Please ensure your .streamlit/secrets.toml file is correctly configured with [openai] and [credentials] sections.")
    st.stop() # Stop the app if secrets is missing

# Initialize OpenAI client (will be initialized after API key check)
client = None 

# --- Session State Initialization ---
# Initialize session state variables if they don't exist
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'current_username' not in st.session_state:
    st.session_state['current_username'] = None
if 'df' not in st.session_state:
    st.session_state['df'] = None
if 'data_summary_text' not in st.session_state: # Text summary for AI
    st.session_state['data_summary_text'] = ""
if 'data_summary_table' not in st.session_state: # Structured data for report table
    st.session_state['data_summary_table'] = []
if 'messages' not in st.session_state:
    st.session_state['messages'] = []
if 'report_content' not in st.session_state:
    st.session_state['report_content'] = [] # List to store report sections (structured)
if 'user_goal' not in st.session_state:
    st.session_state['user_goal'] = "Not specified"
if 'uploaded_file_name' not in st.session_state: # To track if a new file is uploaded
    st.session_state['uploaded_file_name'] = None
if 'openai_client_initialized' not in st.session_state:
    st.session_state['openai_client_initialized'] = False
if 'openai_client' not in st.session_state: # New: Store OpenAI client instance here
    st.session_state['openai_client'] = None
if 'debug_logs' not in st.session_state: # New: For in-app debug logs
    st.session_state['debug_logs'] = []
if 'regression_results' not in st.session_state: # NEW: To store regression summary
    st.session_state['regression_results'] = None
if 'regression_plot_buffer' not in st.session_state: # NEW: To store regression plot buffer
    st.session_state['regression_plot_buffer'] = None
if 'regression_plot_caption' not in st.session_state: # NEW: To store regression plot caption
    st.session_state['regression_plot_caption'] = None


# --- Helper Functions ---

# Function to append debug messages to session state
def append_debug_log(message):
    st.session_state['debug_logs'].append(message)

def check_password():
    """
    Checks if the entered username and password match any of the ones in st.secrets['credentials'].
    This uses plain-text password comparison as per user's request,
    acknowledging the security implications.
    """
    if st.session_state['logged_in']:
        return True

    st.image("SsoLogo.jpg", width=100)
    st.title("Login to Data Preprocessing Assistant")
    st.markdown("---")

    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submit_button = st.form_submit_button("Login")

        if submit_button:
            # Check if the entered username exists in our stored users
            if username in LOGIN_USERS:
                # Compare the entered password with the stored plain-text password
                if password == LOGIN_USERS[username]:
                    st.session_state['logged_in'] = True
                    st.session_state['current_username'] = username # Store current logged-in username
                    st.rerun() # Rerun to switch to the main app
                else:
                    st.error("Invalid username or password.")
            else:
                st.error("Invalid username or password.")
    return False

def check_openai_api_key():
    """
    Attempts a simple OpenAI API call to verify the API key and store the client.
    Returns True if successful, False otherwise.
    """
    if st.session_state['openai_client_initialized'] and st.session_state['openai_client'] is not None:
        return True # Already checked and initialized

    try:
        # Initialize client and store in session state
        st.session_state['openai_client'] = OpenAI(api_key=OPENAI_API_KEY)
        
        # Attempt a simple call to verify the key
        response = st.session_state['openai_client'].chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": "Hello"}],
            max_tokens=5,
            stream=False, # Do not stream for this check
            timeout=10 # Add a timeout for the API call
        )
        st.session_state['openai_client_initialized'] = True
        st.success("OpenAI API key verified successfully! AI features are enabled.")
        return True
    except AuthenticationError:
        st.error("OpenAI API Key is invalid. Please check your .streamlit/secrets.toml file.")
        st.session_state['openai_client'] = None # Clear client on failure
        return False
    except APIConnectionError as e:
        st.error(f"Could not connect to OpenAI API: {e}. Please check your internet connection and firewall settings.")
        st.session_state['openai_client'] = None # Clear client on failure
        return False
    except RateLimitError:
        st.error("OpenAI API rate limit exceeded. Please try again later or check your OpenAI usage.")
        st.session_state['openai_client'] = None # Clear client on failure
        return False
    except requests.exceptions.Timeout:
        st.error("OpenAI API connection timed out during key verification. Please try again.")
        st.session_state['openai_client'] = None
        return False
    except APIStatusError as e: # Catch API specific status errors
        st.error(f"OpenAI API returned an error status: {e.status_code} - {e.response}")
        st.session_state['openai_client'] = None
        return False
    except Exception as e:
        st.error(f"An unexpected error occurred while checking OpenAI API key: {e}")
        st.session_state['openai_client'] = None # Clear client on failure
        return False


def get_data_summary(df):
    """
    Generates a comprehensive summary of the DataFrame's characteristics for AI and report.
    Returns (summary_text, column_details_for_table).
    """
    summary_text = []
    column_details_for_table = [] # For the report table

    summary_text.append(f"Dataset Overview:\n")
    summary_text.append(f"- Number of rows: {df.shape[0]}")
    summary_text.append(f"- Number of columns: {df.shape[1]}")
    summary_text.append(f"- Total duplicate rows: {df.duplicated().sum()}\n")

    column_details_for_table.append(["Column Name", "Data Type", "Missing %", "Stats Summary"]) # Table Header

    # Add a general heading for column details in the text summary
    summary_text.append("Column Details:\n")

    for col in df.columns:
        dtype = df[col].dtype
        missing_percent = df[col].isnull().sum() / len(df) * 100
        
        col_summary_text_parts = []
        col_table_summary = ""

        # For text summary
        col_summary_text_parts.append(f"  - Column '{col}':")
        col_summary_text_parts.append(f"    - Data Type: {dtype}")
        col_summary_text_parts.append(f"    - Missing Values: {missing_percent:.2f}%")

        if pd.api.types.is_numeric_dtype(df[col]):
            desc = df[col].describe()
            col_summary_text_parts.append(f"    - Numerical Stats:")
            col_summary_text_parts.append(f"      - Mean: {desc['mean']:.2f}")
            col_summary_text_parts.append(f"      - Median: {df[col].median():.2f}")
            col_summary_text_parts.append(f"      - Std Dev: {desc['std']:.2f}")
            col_summary_text_parts.append(f"      - Min: {desc['min']:.2f}")
            col_summary_text_parts.append(f"      - Max: {desc['max']:.2f}")
            col_summary_text_parts.append(f"      - 25th Percentile: {desc['25%']:.2f}")
            col_summary_text_parts.append(f"      - 75th Percentile: {desc['75%']:.2f}")
            col_summary_text_parts.append(f"      - Skewness: {df[col].skew():.2f}")
            col_summary_text_parts.append(f"      - Kurtosis: {df[col].kurt():.2f}")
            col_table_summary = (
                f"Mean: {desc['mean']:.2f}, Median: {df[col].median():.2f}, "
                f"Skew: {df[col].skew():.2f}, Missing: {missing_percent:.2f}%"
            )
        elif pd.api.types.is_object_dtype(df[col]) or pd.api.types.is_string_dtype(df[col]):
            unique_count = df[col].nunique()
            top_values = df[col].value_counts().head(3) # Limit to top 3 for table summary
            col_summary_text_parts.append(f"    - Categorical Stats:")
            col_summary_text_parts.append(f"      - Unique Values (Cardinality): {unique_count}")
            if not top_values.empty:
                top_vals_str = ", ".join([f"'{val}': {count}" for val, count in top_values.items()])
                col_summary_text_parts.append(f"      - Top 3 Values and Counts: {top_vals_str}")
                col_table_summary = f"Unique: {unique_count}, Top: {top_vals_str}, Missing: {missing_percent:.2f}%"
            else:
                col_table_summary = f"Unique: {unique_count}, Missing: {missing_percent:.2f}%"
        elif pd.api.types.is_datetime64_any_dtype(df[col]):
            col_summary_text_parts.append(f"    - Date/Time Stats:")
            col_summary_text_parts.append(f"      - Min Date: {df[col].min()}")
            col_summary_text_parts.append(f"      - Max Date: {df[col].max()}")
            col_table_summary = (
                f"Min Date: {df[col].min().strftime('%Y-%m-%d')}, "
                f"Max Date: {df[col].max().strftime('%Y-%m-%d')}, Missing: {missing_percent:.2f}%"
            )
        
        summary_text.append("\n".join(col_summary_text_parts))
        summary_text.append("") # Add a blank line for readability between columns in text summary
        column_details_for_table.append([col, str(dtype), f"{missing_percent:.2f}%", col_table_summary])

    return "\n".join(summary_text), column_details_for_table

def generate_openai_response(prompt, model="gpt-3.5-turbo"):
    """
    Sends a prompt to the OpenAI API and returns the response.
    Explicitly instructs the model NOT to provide Python code snippets or markdown formatting.
    """
    # Retrieve client from session state
    client_instance = st.session_state.get('openai_client')

    if client_instance is None or not st.session_state['openai_client_initialized']:
        append_debug_log("DEBUG: OpenAI client not initialized or missing.") # Debug print
        return "AI features are not enabled due to API key issues. Please check your OpenAI API key."

    try:
        # Filter messages to only include text-based content for the AI API call
        api_messages_history = []
        # Only send the last 5 relevant messages to save tokens and maintain context
        for msg in st.session_state.messages[-5:]:
            if msg["role"] in ["user", "assistant"]: # Only include user and assistant text messages
                api_messages_history.append({"role": msg["role"], "content": msg["content"]})
        
        # Add the current user prompt
        api_messages_history.append({"role": "user", "content": prompt})

        append_debug_log(f"DEBUG: Sending prompt to OpenAI (max_tokens=2000):\n{api_messages_history}\n---") # Debug print
        
        # --- NEW: More robust API call with specific error handling and timeout ---
        try:
            response = client_instance.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a data preprocessing expert. Provide clear, concise, and actionable advice. Do NOT use any markdown formatting (like bolding, italics, code blocks) in your responses. Focus on natural language explanations and interpretations. Always ask the user about their goal for the dataset if not specified, or if a graph is generated, provide an interpretation of that graph. Keep responses concise and to the point."},
                    *api_messages_history
                ],
                temperature=0.7,
                max_tokens=2000,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0,
                timeout=30 # Increased timeout to 30 seconds for API calls
            )
            ai_response_content = response.choices[0].message.content
            append_debug_log(f"DEBUG: Raw OpenAI response received:\n{ai_response_content}\n---") # Debug print
            return ai_response_content
        except AuthenticationError as e:
            append_debug_log(f"DEBUG: API Call AuthenticationError: {e}") # Debug print
            st.error("OpenAI API Key is invalid during chat. Please check your .streamlit/secrets.toml file.")
            return "I'm sorry, my connection to the AI failed due to an invalid API key. Please contact support."
        except APIConnectionError as e:
            append_debug_log(f"DEBUG: API Call ConnectionError: {e}") # Debug print
            st.error(f"Could not connect to OpenAI API during chat: {e}. Please check your internet connection and firewall settings.")
            return "I'm sorry, I'm having trouble connecting to the AI. Please check your internet connection and try again."
        except RateLimitError as e:
            append_debug_log(f"DEBUG: API Call RateLimitError: {e}") # Debug print
            st.error("OpenAI API rate limit exceeded during chat. Please try again in a moment.")
            return "I'm sorry, the AI is experiencing high demand. Please try again in a moment."
        except requests.exceptions.Timeout as e: # Catch specific requests timeout
            append_debug_log(f"DEBUG: API Call Timeout: {e}") # Debug print
            st.error("OpenAI API request timed out. The server took too long to respond. Please try again.")
            return "I'm sorry, the AI took too long to respond. Please try again in a moment."
        except APIStatusError as e: # Catch API specific status errors (e.g., 4xx, 5xx from OpenAI)
            append_debug_log(f"DEBUG: API Call Status Error: {e.status_code} - {e.response}") # Debug print
            st.error(f"OpenAI API returned an error status: {e.status_code}. Please try again later.")
            return f"I'm sorry, the AI encountered an error ({e.status_code}). Please try again later."
        except Exception as e: # Catch any other unexpected errors during the API call
            append_debug_log(f"DEBUG: Unexpected Exception during API call: {e}") # Debug print
            st.error(f"An unexpected error occurred during AI API call: {e}")
            st.exception(e) # Display the full traceback in the Streamlit UI for debugging
            return "I'm sorry, an unexpected error occurred while communicating with the AI. Please try again later."

    except Exception as e: # This outer catch block is for errors *before* the API call (e.g., client not initialized)
        append_debug_log(f"DEBUG: General Exception in generate_openai_response (outer block): {e}") # Debug print
        st.error(f"An unexpected error occurred while generating the AI response: {e}")
        st.exception(e) # Display the full traceback in the Streamlit UI for debugging
        return "I'm sorry, an unexpected error occurred while generating the AI response. Please try again later."


def create_report_doc(report_data, logo_path="SsoLogo.jpg"):
    """
    Generates a Word document report from the accumulated report_data.
    """
    document = Document()

    # Add Logo to the report
    try:
        document.add_picture(logo_path, width=Inches(1.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except FileNotFoundError:
        document.add_paragraph("SSO Consultants Logo (Image not found)")
    except Exception as e:
        document.add_paragraph(f"Error adding logo to report: {e}")

    document.add_heading('Dataset Preprocessing & Analysis Report', level=1)
    document.add_paragraph(f"Date Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"User's Stated Goal: {st.session_state['user_goal']}")
    document.add_paragraph("\n")

    for item in report_data:
        item_type = item.get("type")
        content = item.get("content")

        if item_type == "heading":
            heading = document.add_heading('', level=item.get("level", 2))
            run = heading.add_run(content)
            run.bold = True # Ensure headings are bold
        elif item_type == "text":
            # REVISED LOGIC FOR TEXT CONTENT TO HANDLE NUMBERED LISTS PROPERLY
            lines = content.split('\n')
            for line in lines:
                line = line.strip() # Remove leading/trailing whitespace
                if not line: # Skip empty lines
                    continue

                p = document.add_paragraph()
                # Check if it's a numbered list item
                numbered_list_match = re.match(r'^(\d+\.\s*)(.*)', line)
                if numbered_list_match:
                    # Add the number part (e.g., "1. ")
                    p.add_run(numbered_list_match.group(1))
                    
                    # Get the rest of the line after the number
                    rest_of_line = numbered_list_match.group(2)
                    
                    # This regex is for finding markdown bolding within the rest_of_line
                    parts = re.split(r'(\*\*.*?\*\*)', rest_of_line)
                    for part in parts:
                        if part is None or part == '':
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    # Not a numbered list item, just add as a regular paragraph
                    # Also check for markdown bolding in regular paragraphs
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part is None or part == '':
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

        elif item_type == "table":
            # Original headers and rows from st.session_state['data_summary_table']
            original_headers = item.get("headers", []) # ['Column Name', 'Data Type', 'Missing %', 'Stats Summary']
            original_rows = item.get("rows", []) # [['ID', 'int64', '0.00%', 'Mean: ...'], ...]

            if original_headers and original_rows:
                # Add a specific sub-heading for the table
                table_heading = document.add_heading('', level=3)
                run = table_heading.add_run("Column Details Overview")
                run.bold = True

                # Define new headers for the Word table
                # We want "Column Name\n(Type)", "Missing %", "Stats Summary"
                new_table_headers = ["Column Name\n(Type)", original_headers[2], original_headers[3]]
                
                table = document.add_table(rows=1, cols=len(new_table_headers))
                table.style = 'Table Grid' # Apply a basic table style

                # Add new headers to the Word table
                hdr_cells = table.rows[0].cells
                for i, header_text in enumerate(new_table_headers):
                    hdr_cells[i].text = header_text
                    # Set header bold
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Data type short form mapping
                dtype_map = {
                    'int64': 'num',
                    'float64': 'num',
                    'object': 'str',
                    'category': 'str',
                    'datetime64[ns]': 'date', # Pandas datetime dtype often includes [ns]
                    'datetime64': 'date', # General datetime
                    'bool': 'bool' # Boolean type
                }

                # Add rows, transforming the first two columns
                for row_data in original_rows:
                    # row_data is like ['ID', 'int64', '0.00%', 'Mean: ...']
                    col_name = row_data[0]
                    original_dtype = row_data[1]
                    # Get short form, default to 'other' if not in map
                    short_dtype = dtype_map.get(original_dtype, 'other') 
                    combined_col_info = f"{col_name}\n({short_dtype})"
                    
                    # Create the new row for the Word table
                    new_row_for_table = [combined_col_info, row_data[2], row_data[3]] # Use original missing % and stats summary

                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(new_row_for_table):
                        row_cells[i].text = str(cell_data)

        elif item_type == "image":
            # Image data is expected as BytesIO object
            try:
                document.add_picture(content, width=Inches(6)) # Adjust width as needed
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = item.get("caption", "")
                if caption:
                    document.add_paragraph(caption, style='Caption') # Add caption style if available
            except Exception as e:
                document.add_paragraph(f"Error adding image to report: {e}")
        
        elif item_type == "stat_table": # For structured statistical tables
            table_title = item.get("title")
            df_to_add = item.get("dataframe")

            if table_title and df_to_add is not None:
                document.add_paragraph(table_title, style='Heading 4') # Use a sub-heading for the table
                
                # Create a new table in the document
                table = document.add_table(rows=df_to_add.shape[0] + 1, cols=df_to_add.shape[1])
                table.style = 'Table Grid' # Apply a basic table style

                # Add header row
                for i, col_name in enumerate(df_to_add.columns):
                    cell = table.cell(0, i)
                    cell.text = str(col_name)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Add data rows
                for r_idx, row_data in enumerate(df_to_add.itertuples(index=False), start=1):
                    for c_idx, cell_value in enumerate(row_data):
                        # Format floats to 4 decimal places
                        table.cell(r_idx, c_idx).text = str(f"{cell_value:.4f}" if isinstance(cell_value, (float)) else cell_value)
            
        elif item_type == "regression_summary": # NEW: For regression summary text
            document.add_paragraph(content)

        document.add_paragraph("\n") # Add a blank line after each section for spacing

    # Add a section for the footer in the report
    document.add_section(WD_SECTION_START.NEW_PAGE) # Start new page for footer/disclaimer
    footer_paragraph = document.add_paragraph()
    footer_run = footer_paragraph.add_run("SSO Consultants © 2025 | All Rights Reserved.")
    footer_run.font.size = Pt(9)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document to a BytesIO object
    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0) # Rewind the buffer to the beginning
    return bio

def generate_and_display_graph(df, graph_type, columns):
    """
    Generates a graph (histogram, box plot, scatter plot, correlation heatmap, bar chart)
    and returns its BytesIO buffer and a description.
    """
    img_buffer = io.BytesIO()
    graph_desc = ""
    plt.figure(figsize=(10, 6)) # Set a default figure size

    try:
        if graph_type == "histogram":
            if not columns or not pd.api.types.is_numeric_dtype(df[columns[0]]):
                return None, "Unsupported graph type requested or invalid column for histogram. Please select a numerical column."
            sns.histplot(df[columns[0]].dropna(), kde=True)
            plt.title(f'Histogram of {columns[0]}')
            plt.xlabel(columns[0])
            plt.ylabel('Frequency')
            graph_desc = f"A histogram for the '{columns[0]}' column was generated. It shows the distribution of values for this numerical feature."
        
        elif graph_type == "box_plot":
            if not columns or not pd.api.types.is_numeric_dtype(df[columns[0]]):
                return None, "Unsupported graph type requested or invalid column for box plot. Please select a numerical column."
            sns.boxplot(y=df[columns[0]].dropna())
            plt.title(f'Box Plot of {columns[0]}')
            plt.ylabel(columns[0])
            graph_desc = f"A box plot for the '{columns[0]}' column was generated. It displays the distribution, median, quartiles, and potential outliers of this numerical feature."

        elif graph_type == "scatter_plot":
            if len(columns) != 2 or not pd.api.types.is_numeric_dtype(df[columns[0]]) or not pd.api.types.is_numeric_dtype(df[columns[1]]):
                return None, "Unsupported graph type requested or invalid columns for scatter plot. Please select two numerical columns."
            sns.scatterplot(x=df[columns[0]], y=df[columns[1]])
            plt.title(f'Scatter Plot of {columns[0]} vs {columns[1]}')
            plt.xlabel(columns[0])
            plt.ylabel(columns[1])
            graph_desc = f"A scatter plot of '{columns[0]}' versus '{columns[1]}' was generated. It visualizes the relationship between these two numerical features."

        elif graph_type == "correlation_heatmap":
            numerical_cols = df.select_dtypes(include=['number']).columns
            if numerical_cols.empty:
                return None, "No numerical columns found to generate a correlation heatmap."
            correlation_matrix = df[numerical_cols].corr()
            sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', fmt=".2f")
            plt.title("Correlation Matrix of Numerical Features")
            graph_desc = "A correlation heatmap of all numerical features was generated. It displays the pairwise correlation coefficients, indicating the strength and direction of linear relationships between variables."

        elif graph_type == "bar_chart":
            if not columns or not (pd.api.types.is_object_dtype(df[columns[0]]) or pd.api.types.is_string_dtype(df[columns[0]]) or pd.api.types.is_categorical_dtype(df[columns[0]])):
                return None, "Unsupported graph type requested or invalid column for bar chart. Please select a categorical column."
            
            # For bar charts, limit to top 10 categories to avoid clutter
            value_counts = df[columns[0]].value_counts().head(10)
            if value_counts.empty:
                return None, f"No data found for bar chart in column '{columns[0]}'."

            sns.barplot(x=value_counts.index, y=value_counts.values)
            plt.title(f'Frequency of {columns[0]} (Top {len(value_counts)})')
            plt.xlabel(columns[0])
            plt.ylabel('Count')
            plt.xticks(rotation=45, ha='right') # Rotate labels for readability
            plt.tight_layout() # Adjust layout to prevent labels overlapping
            graph_desc = f"A bar chart showing the frequency of top categories in '{columns[0]}' was generated. It helps visualize the distribution of categorical values."
        
        else:
            return None, "Unsupported graph type requested. Please ask for a histogram, box plot, scatter plot, correlation heatmap, or bar chart."

        plt.savefig(img_buffer, format='png', bbox_inches='tight') # Save to buffer
        img_buffer.seek(0) # Rewind the buffer
        plt.close() # Close the plot to free memory
        return img_buffer, graph_desc

    except Exception as e:
        plt.close() # Ensure plot is closed even on error
        return None, f"An error occurred while generating the graph: {e}. Please check your column selections and data."


def perform_statistical_test(df, test_type, col1=None, col2=None):
    """
    Performs the selected statistical test and returns the results as a formatted string,
    structured DataFrame(s), an error message, and optionally a plot buffer and caption.
    
    Returns (results_str, structured_results_for_ui, error_message, plot_buffer, plot_caption).
    structured_results_for_ui can be a DataFrame, tuple of DataFrames, or None.
    plot_buffer and plot_caption are for tests that generate plots (e.g., future tests).
    """
    results_str = ""
    structured_results_for_ui = None
    error_message = None
    plot_buffer = None
    plot_caption = None

    try:
        # NEW: Cronbach's Alpha Test
        if test_type == "Cronbach’s Alpha (Reliability)":
            # For Cronbach's Alpha, col1 is expected to be a list of column names (items)
            if not (isinstance(col1, list) and len(col1) >= 2):
                error_message = "Cronbach’s Alpha: Please select at least two numerical columns (items) for the test."
            else:
                selected_columns = col1
                numeric_cols_data = df[selected_columns].dropna()

                for col in selected_columns:
                    if not pd.api.types.is_numeric_dtype(df[col]):
                        error_message = f"Cronbach’s Alpha: All selected columns must be numerical. '{col}' is not."
                        break # Exit loop if non-numerical column found

                if not error_message and (numeric_cols_data.empty or len(numeric_cols_data) < 2):
                    error_message = "Cronbach’s Alpha: Not enough valid data points after dropping NaNs (requires at least 2 rows)."
                
                if not error_message:
                    alpha, n_items = pg.cronbach_alpha(data=numeric_cols_data, return_N=True)

                    results_df = pd.DataFrame({
                        "Metric": ["Alpha Value", "Number of Items"],
                        "Value": [alpha, n_items]
                    })
                    
                    results_str = (
                        f"Cronbach’s Alpha (Reliability) Test Results:\n"
                        f"  Alpha Value: {alpha:.4f}\n"
                        f"  Number of Items: {n_items}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = results_df
        
        # NEW: Z-Test (Two Sample Means)
        elif test_type == "Z-Test (Two Sample Means)":
            append_debug_log(f"DEBUG Z-Test: col1={col1}, col2={col2}")
            if not pd.api.types.is_numeric_dtype(df[col1]):
                error_message = f"Z-Test: Numerical variable '{col1}' must be numerical."
            elif not (pd.api.types.is_object_dtype(df[col2]) or pd.api.types.is_string_dtype(df[col2]) or pd.api.types.is_categorical_dtype(df[col2])):
                error_message = f"Z-Test: Grouping variable '{col2}' must be categorical."
            else:
                unique_groups = df[col2].dropna().unique()
                if len(unique_groups) != 2:
                    error_message = f"Z-Test: Grouping variable '{col2}' must have exactly 2 distinct groups. Found {len(unique_groups)}."
                else:
                    group1_name, group2_name = unique_groups[0], unique_groups[1]
                    group1_data = df[col1][df[col2] == group1_name].dropna()
                    group2_data = df[col1][df[col2] == group2_name].dropna()

                    if len(group1_data) < 30 or len(group2_data) < 30: # Z-test assumes large sample sizes (n >= 30) or known population std dev
                        st.warning("Warning: Z-Test is typically used for large sample sizes (n >= 30) or when population standard deviation is known. Consider a T-Test if sample sizes are small and population standard deviations are unknown.")
                        # This is a warning, not an error that stops the test. The user still wants to run it.

                    # Calculate sample statistics
                    n1, mean1, std1 = len(group1_data), group1_data.mean(), group1_data.std()
                    n2, mean2, std2 = len(group2_data), group2_data.mean(), group2_data.std()

                    if std1 == 0 and std2 == 0:
                        error_message = "Z-Test: Both groups have zero variance, cannot perform test."
                    elif n1 == 0 or n2 == 0:
                        error_message = f"Z-Test: One or both groups have no data for '{col1}' after dropping NaNs."
                    else:
                        # Calculate pooled standard deviation if assuming equal variance (similar to T-test, but for Z-test context)
                        # For a standard Z-test with unknown population SD, we use sample SD (large n assumption)
                        # Standard Error of the Difference
                        se_diff = np.sqrt((std1**2 / n1) + (std2**2 / n2))

                        if se_diff == 0:
                             error_message = "Z-Test: Standard error of the difference is zero. Cannot perform test."
                        else:
                            z_statistic = (mean1 - mean2) / se_diff
                            p_value = 2 * (1 - stats.norm.cdf(abs(z_statistic))) # Two-tailed p-value

                            group_stats_data = {
                                'Group': [group1_name, group2_name],
                                'N': [n1, n2],
                                'Mean': [mean1, mean2],
                                'Std. Deviation': [std1, std2]
                            }
                            group_stats_df = pd.DataFrame(group_stats_data)

                            test_results_data = {
                                'Statistic': ['Z-statistic', 'P-value'],
                                'Value': [z_statistic, p_value]
                            }
                            test_results_df = pd.DataFrame(test_results_data)

                            results_str = (
                                f"Z-Test (Two Sample Means) Results for '{col1}' by '{col2}' ({group1_name} vs {group2_name}):\n"
                                f"  Z-statistic: {z_statistic:.4f}\n"
                                f"  P-value: {p_value:.4f}\n"
                                "Interpretation will be provided by the AI."
                            )
                            structured_results_for_ui = (group_stats_df, test_results_df)

        elif test_type == "anova":
            append_debug_log(f"DEBUG ANOVA: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG ANOVA: is_numeric_dtype(df[{col1}])={pd.api.types.is_numeric_dtype(df[col1])}")
            append_debug_log(f"DEBUG ANOVA: is_categorical_dtype(df[{col2}])={pd.api.types.is_categorical_dtype(df[col2])} | is_object_dtype(df[{col2}])={pd.api.types.is_object_dtype(df[col2])} | is_string_dtype(df[{col2}])={pd.api.types.is_string_dtype(df[col2])}")
            if not pd.api.types.is_numeric_dtype(df[col1]):
                error_message = f"ANOVA: Dependent variable '{col1}' must be numerical."
            elif not (pd.api.types.is_object_dtype(df[col2]) or pd.api.types.is_string_dtype(df[col2]) or pd.api.types.is_categorical_dtype(df[col2])):
                error_message = f"ANOVA: Independent variable '{col2}' must be categorical."
            else:
                # Drop NaNs from both columns for consistent analysis
                clean_df = df[[col1, col2]].dropna()
                
                groups = [clean_df[col1][clean_df[col2] == g] for g in clean_df[col2].unique()]
                append_debug_log(f"DEBUG ANOVA: unique_groups={clean_df[col2].unique()}, len(groups)={len(groups)}")
                
                if len(groups) < 2:
                    error_message = f"ANOVA: Independent variable '{col2}' needs at least 2 distinct groups."
                elif any(len(g) == 0 for g in groups):
                    error_message = f"ANOVA: Some groups in '{col2}' have no data for '{col1}' after dropping NaNs."
                else:
                    # Perform ANOVA using scipy
                    f_statistic_scipy, p_value_scipy = stats.f_oneway(*groups)

                    # Calculate Sum of Squares (SS) and Degrees of Freedom (df) for ANOVA table
                    grand_mean = clean_df[col1].mean()
                    
                    # Sum of Squares Total (SST)
                    sst = np.sum((clean_df[col1] - grand_mean)**2)
                    df_total = len(clean_df) - 1

                    # Sum of Squares Between (SSB)
                    ssb = 0
                    for g in clean_df[col2].unique():
                        group_data = clean_df[col1][clean_df[col2] == g]
                        ssb += len(group_data) * (group_data.mean() - grand_mean)**2
                    df_between = len(clean_df[col2].unique()) - 1

                    # Sum of Squares Within (SSW)
                    ssw = np.sum((clean_df[col1] - clean_df.groupby(col2)[col1].transform('mean'))**2)
                    df_within = len(clean_df) - len(clean_df[col2].unique())

                    # Mean Squares
                    msb = ssb / df_between if df_between > 0 else np.nan
                    msw = ssw / df_within if df_within > 0 else np.nan

                    # F-statistic (re-calculated to match SS/MS for table consistency, should be close to scipy's)
                    f_stat_calculated = msb / msw if msw > 0 else np.nan

                    # Create ANOVA Summary DataFrame
                    anova_summary_data = {
                        'Source of Variation': ['Between Groups', 'Within Groups', 'Total'],
                        'Sum of Squares (SS)': [ssb, ssw, sst],
                        'df': [df_between, df_within, df_total],
                        'Mean Squares (MS)': [msb, msw, np.nan], # MS for Total is not typically reported
                        'F': [f_stat_calculated, np.nan, np.nan], # F-stat only for Between Groups
                        'P-value': [p_value_scipy, np.nan, np.nan] # P-value only for Between Groups
                    }
                    anova_df = pd.DataFrame(anova_summary_data)
                    
                    results_str = (
                        f"ANOVA Test Results for '{col1}' by '{col2}':\n"
                        f"  F-statistic: {f_statistic_scipy:.4f}\n" # Use scipy's F-stat for the initial text summary
                        f"  P-value: {p_value_scipy:.4f}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = anova_df # Return the single DataFrame

        elif test_type == "independent_t_test":
            append_debug_log(f"DEBUG Independent T-test: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG Independent T-test: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]):
                error_message = f"Independent T-test: Numerical variable '{col1}' must be numerical."
            elif not (pd.api.types.is_object_dtype(df[col2]) or pd.api.types.is_string_dtype(df[col2]) or pd.api.types.is_categorical_dtype(df[col2])):
                error_message = f"Independent T-test: Grouping variable '{col2}' must be categorical."
            else:
                unique_groups = df[col2].unique()
                append_debug_log(f"DEBUG Independent T-test: unique_groups={unique_groups}, len(unique_groups)={len(unique_groups)}")
                if len(unique_groups) != 2:
                    error_message = f"Independent T-test: Grouping variable '{col2}' must have exactly 2 distinct groups. Found {len(unique_groups)}."
                else:
                    group1_name = unique_groups[0]
                    group2_name = unique_groups[1]
                    group1_data = df[col1][df[col2] == group1_name].dropna()
                    group2_data = df[col1][df[col2] == group2_name].dropna()
                    append_debug_log(f"DEBUG Independent T-test: group1_data_len={len(group1_data)}, group2_data_len={len(group2_data)}")
                    if len(group1_data) == 0 or len(group2_data) == 0:
                        error_message = f"Independent T-test: One or both groups have no data for '{col1}' after dropping NaNs."
                    else:
                        t_statistic, p_value = stats.ttest_ind(group1_data, group2_data, equal_var=True) # Assuming equal variances for this specific test
                        
                        # Prepare structured results (Excel-like output)
                        group_stats_data = {
                            'Group': [group1_name, group2_name],
                            'N': [len(group1_data), len(group2_data)],
                            'Mean': [group1_data.mean(), group2_data.mean()],
                            'Std. Deviation': [group1_data.std(), group2_data.std()]
                        }
                        group_stats_df = pd.DataFrame(group_stats_data)

                        test_results_data = {
                            'Statistic': ['T-statistic', 'P-value'],
                            'Value': [t_statistic, p_value]
                        }
                        test_results_df = pd.DataFrame(test_results_data)

                        results_str = (
                            f"Independent T-test (Equal Variances Assumed) Results for '{col1}' by '{col2}' ({group1_name} vs {group2_name}):\n"
                            f"  T-statistic: {t_statistic:.4f}\n"
                            f"  P-value: {p_value:.4f}\n"
                            "Interpretation will be provided by the AI."
                        )
                        # Store both dataframes in a tuple for structured_results_for_ui
                        structured_results_for_ui = (group_stats_df, test_results_df)

        elif test_type == "chi_squared_test":
            append_debug_log(f"DEBUG Chi-squared: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG Chi-squared: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not (pd.api.types.is_object_dtype(df[col1]) or pd.api.types.is_string_dtype(df[col1]) or pd.api.types.is_categorical_dtype(df[col1])):
                error_message = f"Chi-squared: Column 1 '{col1}' must be categorical."
            elif not (pd.api.types.is_object_dtype(df[col2]) or pd.api.types.is_string_dtype(df[col2]) or pd.api.types.is_categorical_dtype(df[col2])):
                error_message = f"Chi-squared: Column 2 '{col2}' must be categorical."
            else:
                contingency_table = pd.crosstab(df[col1], df[col2])
                append_debug_log(f"DEBUG Chi-squared: contingency_table_shape={contingency_table.shape}, sum={contingency_table.sum().sum()}")
                if contingency_table.empty or contingency_table.sum().sum() == 0:
                     error_message = f"Chi-squared: No valid data to form a contingency table for '{col1}' and '{col2}'."
                else:
                    chi2, p_value, dof, expected = stats.chi2_contingency(contingency_table)
                    
                    # Create DataFrames for Observed and Expected tables
                    observed_df = contingency_table
                    expected_df = pd.DataFrame(expected, index=contingency_table.index, columns=contingency_table.columns)

                    results_str = (
                        f"Chi-squared Test Results for '{col1}' and '{col2}':\n"
                        f"  Chi-squared statistic: {chi2:.4f}\n"
                        f"  P-value: {p_value:.4f}\n"
                        f"  Degrees of Freedom (dof): {dof}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = (observed_df, expected_df, chi2, p_value, dof) # Return tuple of DFs and key stats

        elif test_type == "paired_t_test":
            append_debug_log(f"DEBUG Paired T-test: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG Paired T-test: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]) or not pd.api.types.is_numeric_dtype(df[col2]):
                error_message = f"Paired T-test: Both '{col1}' and '{col2}' must be numerical variables."
            else:
                # Drop rows where either of the two columns has NaN to ensure paired data
                paired_df = df[[col1, col2]].dropna()
                if len(paired_df) < 2: # Need at least 2 pairs to run t-test
                    error_message = f"Paired T-test: Not enough paired data points after dropping NaNs. Need at least 2."
                else:
                    t_statistic, p_value = stats.ttest_rel(paired_df[col1], paired_df[col2])

                    # Calculate means and std devs for display
                    mean1 = paired_df[col1].mean()
                    std1 = paired_df[col1].std()
                    mean2 = paired_df[col2].mean()
                    std2 = paired_df[col2].std()
                    n_pairs = len(paired_df)

                    # Create Group Statistics table for paired data
                    group_stats_data = {
                        'Variable': [col1, col2],
                        'N': [n_pairs, n_pairs],
                        'Mean': [mean1, mean2],
                        'Std. Deviation': [std1, std2]
                    }
                    group_stats_df = pd.DataFrame(group_stats_data)

                    # Create Test Results table
                    test_results_data = {
                        'Statistic': ['T-statistic', 'P-value'],
                        'Value': [t_statistic, p_value]
                    }
                    test_results_df = pd.DataFrame(test_results_data)

                    results_str = (
                        f"Paired T-test Results for '{col1}' and '{col2}':\n"
                        f"  T-statistic: {t_statistic:.4f}\n"
                        f"  P-value: {p_value:.4f}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = (group_stats_df, test_results_df)

        elif test_type == "pearson_correlation (Validity-Linear)":
            append_debug_log(f"DEBUG Pearson Correlation: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG Pearson Correlation: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]) or not pd.api.types.is_numeric_dtype(df[col2]):
                error_message = f"Pearson Correlation: Both '{col1}' and '{col2}' must be numerical variables."
            else:
                # Drop rows where either of the two columns has NaN to ensure consistent length
                clean_df = df[[col1, col2]].dropna()
                if len(clean_df) < 2: # Need at least 2 data points for correlation
                    error_message = f"Pearson Correlation: Not enough data points after dropping NaNs. Need at least 2."
                else:
                    correlation_coefficient, p_value = stats.pearsonr(clean_df[col1], clean_df[col2])
                    covariance = clean_df[col1].cov(clean_df[col2]) # Calculate covariance
                    n_obs = len(clean_df)

                    # Create Correlation Results table
                    corr_results_data = {
                        'Statistic': ['Pearson Correlation (r)', 'Covariance', 'P-value', 'N'],
                        'Value': [correlation_coefficient, covariance, p_value, n_obs]
                    }
                    corr_results_df = pd.DataFrame(corr_results_data)

                    results_str = (
                        f"Pearson Correlation Results for '{col1}' and '{col2}':\n"
                        f"  Correlation Coefficient (r): {correlation_coefficient:.4f}\n"
                        f"  P-value: {p_value:.4f}\n"
                        f"  Covariance: {covariance:.4f}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = corr_results_df

        elif test_type == "spearman_rank_correlation":
            append_debug_log(f"DEBUG Spearman Correlation: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG Spearman Correlation: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]) or not pd.api.types.is_numeric_dtype(df[col2]): # Spearman can also use ordinal, but for simplicity, we'll stick to numerical selection
                error_message = f"Spearman Rank Correlation: Both '{col1}' and '{col2}' must be numerical variables."
            else:
                # Drop rows where either of the two columns has NaN to ensure consistent length
                clean_df = df[[col1, col2]].dropna()
                if len(clean_df) < 2: # Need at least 2 data points for correlation
                    error_message = f"Spearman Rank Correlation: Not enough data points after dropping NaNs. Need at least 2."
                else:
                    correlation_coefficient, p_value = stats.spearmanr(clean_df[col1], clean_df[col2])
                    n_obs = len(clean_df)

                    # Create Correlation Results table
                    spearman_results_data = {
                        'Statistic': ['Spearman Correlation (rho)', 'P-value', 'N'],
                        'Value': [correlation_coefficient, p_value, n_obs]
                    }
                    spearman_results_df = pd.DataFrame(spearman_results_data)

                    results_str = (
                        f"Spearman Rank Correlation Results for '{col1}' and '{col2}':\n"
                        f"  Correlation Coefficient (rho): {correlation_coefficient:.4f}\n"
                        f"  P-value: {p_value:.4f}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = spearman_results_df

        elif test_type == "f_test_two_sample_for_variances":
            append_debug_log(f"DEBUG F-Test for Variances: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG F-Test for Variances: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]) or not pd.api.types.is_numeric_dtype(df[col2]):
                error_message = f"F-Test for Variances: Both '{col1}' and '{col2}' must be numerical variables."
            else:
                data1 = df[col1].dropna()
                data2 = df[col2].dropna()
                if len(data1) < 2 or len(data2) < 2:
                    error_message = f"F-Test for Variances: Both columns must have at least 2 non-null data points. '{col1}' has {len(data1)}, '{col2}' has {len(data2)}."
                else:
                    var1 = np.var(data1, ddof=1) # Sample variance
                    var2 = np.var(data2, ddof=1) # Sample variance

                    if var1 == 0 and var2 == 0:
                        error_message = "F-Test for Variances: Both variances are zero. Cannot perform test."
                    elif var1 == 0:
                        f_statistic = np.inf # If one variance is zero, F-stat is infinite
                    elif var2 == 0:
                        f_statistic = 0.0 # If other variance is zero, F-stat is zero
                    else:
                        f_statistic = var1 / var2 # F-statistic can be < 1

                    df1 = len(data1) - 1
                    df2 = len(data2) - 1

                    if f_statistic == np.inf:
                        p_value = 0.0 # Very small p-value
                    elif f_statistic == 0.0: # Check for 0.0 explicitly
                        p_value = 1.0 # Very large p-value
                    else:
                        # For two-tailed test, p-value is 2 * min(cdf(F), 1 - cdf(F))
                        # The degrees of freedom order depends on which variance is in the numerator
                        p_value = 2 * min(stats.f.cdf(f_statistic, df1, df2), 1 - stats.f.cdf(f_statistic, df1, df2))


                    f_test_data = {
                        'Statistic': ['Variance 1', 'Observations 1', 'df 1', 'Variance 2', 'Observations 2', 'df 2', 'F-statistic', 'P-value'],
                        'Value': [var1, float(len(data1)), float(df1), var2, float(len(data2)), float(df2), f_statistic, p_value]
                    }
                    f_test_df = pd.DataFrame(f_test_data)

                    results_str = (
                        f"F-Test Two-Sample for Variances Results for '{col1}' and '{col2}':\n"
                        f"  F-statistic: {f_statistic:.4f}\n"
                        f"  P-value: {p_value:.4f}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = f_test_df

        elif test_type == "t_test_two_sample_assuming_unequal_variances":
            append_debug_log(f"DEBUG Unequal Variances T-test: col1={col1}, col2={col2}")
            append_debug_log(f"DEBUG Unequal Variances T-test: df[{col1}].dtype={df[col1].dtype}, df[{col2}].dtype={df[col2].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]):
                error_message = f"T-test (Unequal Variances): Numerical variable '{col1}' must be numerical."
            elif not (pd.api.types.is_object_dtype(df[col2]) or pd.api.types.is_string_dtype(df[col2]) or pd.api.types.is_categorical_dtype(df[col2])):
                error_message = f"T-test (Unequal Variances): Grouping variable '{col2}' must be categorical."
            else:
                unique_groups = df[col2].unique()
                append_debug_log(f"DEBUG Unequal Variances T-test: unique_groups={unique_groups}, len(unique_groups)={len(unique_groups)}")
                if len(unique_groups) != 2:
                    error_message = f"T-test (Unequal Variances): Grouping variable '{col2}' must have exactly 2 distinct groups. Found {len(unique_groups)}."
                else:
                    group1_name = unique_groups[0]
                    group2_name = unique_groups[1]
                    group1_data = df[col1][df[col2] == group1_name].dropna()
                    group2_data = df[col1][df[col2] == group2_name].dropna()
                    append_debug_log(f"DEBUG Unequal Variances T-test: group1_data_len={len(group1_data)}, group2_data_len={len(group2_data)}")
                    if len(group1_data) == 0 or len(group2_data) == 0:
                        error_message = f"T-test (Unequal Variances): One or both groups have no data for '{col1}' after dropping NaNs."
                    else:
                        t_statistic, p_value = stats.ttest_ind(group1_data, group2_data, equal_var=False) # Assuming unequal variances
                        
                        # Prepare structured results (Excel-like output)
                        group_stats_data = {
                            'Group': [group1_name, group2_name],
                            'N': [len(group1_data), len(group2_data)],
                            'Mean': [group1_data.mean(), group2_data.mean()],
                            'Std. Deviation': [group1_data.std(), group2_data.std()]
                        }
                        group_stats_df = pd.DataFrame(group_stats_data)

                        test_results_data = {
                            'Statistic': ['T-statistic', 'P-value'],
                            'Value': [t_statistic, p_value]
                        }
                        test_results_df = pd.DataFrame(test_results_data)

                        results_str = (
                            f"Independent T-test (Unequal Variances Assumed) Results for '{col1}' by '{col2}' ({group1_name} vs {group2_name}):\n"
                            f"  T-statistic: {t_statistic:.4f}\n"
                            f"  P-value: {p_value:.4f}\n"
                            "Interpretation will be provided by the AI."
                        )
                        structured_results_for_ui = (group_stats_df, test_results_df)

        elif test_type == "shapiro_wilk_test_normality": # FIXED: Corrected name to match selected_test conversion
            append_debug_log(f"DEBUG Shapiro-Wilk Test: col1={col1}") # col1 here is actually stat_col_single
            append_debug_log(f"DEBUG Shapiro-Wilk Test: df[{col1}].dtype={df[col1].dtype}")
            if not pd.api.types.is_numeric_dtype(df[col1]):
                error_message = f"Shapiro-Wilk Test: Variable '{col1}' must be numerical."
            else:
                data = df[col1].dropna()
                if len(data) < 3: # Shapiro-Wilk requires at least 3 data points
                    error_message = f"Shapiro-Wilk Test: Not enough data points after dropping NaNs. Need at least 3. Found {len(data)}."
                else:
                    shapiro_w, shapiro_p = stats.shapiro(data)
                    
                    normality_results_data = {
                        'Statistic': ['W-statistic', 'P-value', 'N'],
                        'Value': [shapiro_w, shapiro_p, float(len(data))] # Ensure N is float for consistent formatting
                    }
                    normality_results_df = pd.DataFrame(normality_results_data)

                    results_str = (
                        f"Shapiro-Wilk Test for Normality on '{col1}':\n"
                        f"  W-statistic: {shapiro_w:.4f}\n"
                        f"  P-value: {shapiro_p:.4f}\n"
                        "Interpretation will be provided by the AI."
                    )
                    structured_results_for_ui = normality_results_df

        else:
            error_message = "Unsupported statistical test type selected."

    except Exception as e:
        error_message = f"An error occurred during the statistical test: {e}. Please check your column selections and data."
        st.exception(e) # Display traceback in UI for debugging
        append_debug_log(f"DEBUG: Exception in perform_statistical_test: {e}") # Debug print

    return results_str, structured_results_for_ui, error_message, plot_buffer, plot_caption # Return plot_buffer and plot_caption (will be None for most stat tests)


def perform_linear_regression(df, dependent_var, independent_var):
    """
    Performs simple linear regression and returns summary, plot buffer, and caption.
    Returns (summary_text, plot_buffer, plot_caption, error_message)
    """
    summary_text = None
    plot_buffer = io.BytesIO()
    plot_caption = None
    error_message = None

    if not pd.api.types.is_numeric_dtype(df[dependent_var]):
        error_message = f"Dependent variable '{dependent_var}' must be numerical."
    elif not pd.api.types.is_numeric_dtype(df[independent_var]):
        error_message = f"Independent variable '{independent_var}' must be numerical."
    else:
        try:
            # Drop NaNs for the selected columns
            temp_df = df[[dependent_var, independent_var]].dropna()
            
            if len(temp_df) < 2:
                error_message = "Not enough valid data points for regression after dropping NaNs (requires at least 2)."
            else:
                # Add a constant to the independent variable for statsmodels
                X = sm.add_constant(temp_df[independent_var])
                y = temp_df[dependent_var]

                # Create and fit the OLS model
                model = sm.OLS(y, X)
                results = model.fit()

                summary_text = results.summary().as_text() # Get the full text summary

                # Generate the regression plot
                plt.figure(figsize=(10, 6))
                sns.regplot(x=temp_df[independent_var], y=temp_df[dependent_var], scatter_kws={'alpha':0.3})
                plt.title(f'Regression Plot: {dependent_var} vs {independent_var}')
                plt.xlabel(independent_var)
                plt.ylabel(dependent_var)
                plot_caption = f"Simple Linear Regression Plot of '{dependent_var}' against '{independent_var}'."
                
                plt.savefig(plot_buffer, format='png', bbox_inches='tight')
                plot_buffer.seek(0)
                plt.close()

        except Exception as e:
            plt.close() # Ensure plot is closed even on error
            error_message = f"An error occurred during linear regression: {e}. Please check your column selections and data."
            st.exception(e)

    return summary_text, plot_buffer, plot_caption, error_message


# --- Main Application Logic ---
def main_app():
    st.set_page_config(layout="wide", page_title="SSO Data Preprocessing Assistant")

    # --- Safe Logout ---
    if st.session_state.get('logged_in', False):
        with st.sidebar:
            st.markdown("---")
            if st.button("🚪 Logout"):
                keys_to_clear = [
                    'logged_in', 'current_username', 'df', 'data_summary_text', 'data_summary_table',
                    'messages', 'report_content', 'user_goal', 'uploaded_file_name',
                    'openai_client_initialized', 'openai_client', 'debug_logs',
                    'regression_results', 'regression_plot_buffer', 'regression_plot_caption'
                ]
                for key in keys_to_clear:
                    if key in st.session_state:
                        del st.session_state[key]
                st.success("You have been logged out.")
                st.rerun()

    # Display Logo at the top of the main app
    st.image("SsoLogo.jpg", width=100) # Adjust width as needed
    st.title("Data Preprocessing Assistant")
    st.write(f"Welcome, {st.session_state.get('current_username', 'User')}!")

    # --- OpenAI API Key Check ---
    # Perform this check once after login
    if not st.session_state['openai_client_initialized'] or st.session_state['openai_client'] is None:
        with st.spinner("Verifying OpenAI API key..."):
            if not check_openai_api_key():
                st.stop() # Stop the app if API key is invalid or connection fails

    st.sidebar.header("Upload Dataset")
    uploaded_file = st.sidebar.file_uploader("Choose a CSV or Excel file", type=["csv", "xlsx"])

    if uploaded_file is not None:
        # Check if a new file is uploaded or if df is not yet loaded
        if st.session_state['df'] is None or uploaded_file.name != st.session_state.get('uploaded_file_name'):
            st.session_state['messages'] = [] # Clear chat history for new file
            st.session_state['report_content'] = [] # Clear report content for new file
            st.session_state['user_goal'] = "Not specified" # Reset user goal
            st.session_state['uploaded_file_name'] = uploaded_file.name # Store file name to detect new upload
            st.session_state['debug_logs'] = [] # Clear debug logs for new file
            st.session_state['regression_results'] = None # Reset regression results
            st.session_state['regression_plot_buffer'] = None # Reset regression plot
            st.session_state['regression_plot_caption'] = None # Reset regression plot caption


            try:
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                elif uploaded_file.name.endswith('.xlsx'):
                    df = pd.read_excel(uploaded_file)
                st.session_state['df'] = df

                # Generate data summary and store it
                summary_text, summary_table = get_data_summary(df)
                st.session_state['data_summary_text'] = summary_text
                st.session_state['data_summary_table'] = summary_table

                # Add dataset overview to report content
                st.session_state['report_content'].append({"type": "heading", "level": 2, "content": "Dataset Overview"})
                # Split the text summary to only include the general overview, not column details
                overview_text_end_index = summary_text.find("Column Details:")
                if overview_text_end_index != -1:
                    overview_text = summary_text[:overview_text_end_index].strip()
                else:
                    overview_text = summary_text.strip()
                st.session_state['report_content'].append({"type": "text", "content": overview_text})
                st.session_state['report_content'].append({"type": "table", "headers": summary_table[0], "rows": summary_table[1:]}) # Column details table

                # Initial prompt to OpenAI with data summary
                initial_ai_prompt = (
                    "Here is a detailed summary of the user's dataset:\n\n"
                    f"{st.session_state['data_summary_text']}\n\n"
                    "Based on this, what are the initial preprocessing considerations? "
                    "Please also ask the user about their primary goal (e.g., classification, regression, exploratory analysis) for this dataset."
                )
                append_debug_log(f"DEBUG: Initial AI prompt:\n{initial_ai_prompt}\n---") # Debug print
                with st.spinner("Analyzing data and generating initial insights..."):
                    initial_response = generate_openai_response(initial_ai_prompt)
                    st.session_state.messages.append({"role": "assistant", "content": initial_response})
                    st.session_state.report_content.append({"type": "heading", "level": 2, "content": "Initial Preprocessing Considerations"})
                    st.session_state.report_content.append({"type": "text", "content": initial_response})

            except Exception as e:
                st.error(f"Error reading file: {e}. Please ensure it's a valid CSV or Excel file.")
                st.session_state['df'] = None # Reset df on error
                st.stop() # Stop further execution if file reading fails

    if st.session_state['df'] is not None:
        # --- Persistent Data Preview Area ---
        st.subheader("Current Dataset Preview (Top 5 Rows):")
        
        # Data type short form mapping (re-used from report generation)
        dtype_map = {
            'int64': 'int', # Changed from 'num' to 'int' for clarity in UI
            'float64': 'float', # Changed from 'num' to 'float' for clarity in UI
            'object': 'str',
            'category': 'str',
            'datetime64[ns]': 'date',
            'datetime64': 'date',
            'bool': 'bool' # Boolean type
        }

        # Create a copy of the DataFrame to modify column names for display
        display_df = st.session_state['df'].head().copy()
        new_columns = []
        for col in display_df.columns:
            original_dtype = str(st.session_state['df'][col].dtype) # Get original dtype from full df
            short_dtype = dtype_map.get(original_dtype, 'other')
            new_columns.append(f"{col} ({short_dtype})")
        display_df.columns = new_columns

        st.dataframe(display_df)
        st.write(f"Shape: {st.session_state['df'].shape[0]} rows, {st.session_state['df'].shape[1]} columns")
        st.markdown("---") # Separator for clarity

        st.subheader("Chat with your Data Preprocessing Assistant")

        # Display chat messages from history with different colors
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                if message["role"] == "graph" and "content" in message:
                    st.image(message["content"], caption=message.get("caption", ""), use_container_width=True)
                elif message["role"] == "dataframe": # Handle dataframe messages
                    st.subheader(message.get("title", "Statistical Table")) # Display title for the table
                    st.dataframe(message["content"])
                elif message["role"] == "code": # Handle regression summary (plain text, but styled as code for readability)
                    st.subheader(message.get("title", "Regression Summary"))
                    st.code(message["content"], language='text')
                else:
                    st.markdown(message["content"])

        # --- Interactive Data Exploration Controls (New Section) ---
        st.sidebar.markdown("---")
        st.sidebar.header("Generate Specific Graphs")
        
        graph_options = ["Select a graph type", "Histogram", "Box Plot", "Scatter Plot", "Correlation Heatmap", "Bar Chart"]
        selected_graph_type = st.sidebar.selectbox("Choose Graph Type:", graph_options, key="graph_type_select")

        all_columns = st.session_state['df'].columns.tolist()
        numerical_columns = st.session_state['df'].select_dtypes(include=['number']).columns.tolist()
        categorical_columns = st.session_state['df'].select_dtypes(include=['object', 'category']).columns.tolist()

        columns_to_plot = []
        if selected_graph_type in ["Histogram", "Box Plot"]:
            col_options = numerical_columns
            selected_col = st.sidebar.selectbox(f"Select Column for {selected_graph_type}:", ["Select a column"] + col_options, key="single_col_select")
            if selected_col != "Select a column":
                columns_to_plot = [selected_col]
        elif selected_graph_type == "Bar Chart":
            col_options = categorical_columns
            selected_col = st.sidebar.selectbox(f"Select Column for {selected_graph_type}:", ["Select a column"] + col_options, key="single_col_select_cat")
            if selected_col != "Select a column":
                columns_to_plot = [selected_col]
        elif selected_graph_type == "Scatter Plot":
            col1 = st.sidebar.selectbox("Select X-axis Column:", ["Select X"] + numerical_columns, key="scatter_x_select")
            col2 = st.sidebar.selectbox("Select Y-axis Column:", ["Select Y"] + numerical_columns, key="scatter_y_select")
            if col1 != "Select X" and col2 != "Select Y":
                columns_to_plot = [col1, col2]
        # Correlation Heatmap doesn't need column selection here, it uses all numerical

        if st.sidebar.button(f"Generate {selected_graph_type} Chart"): # Changed button label for clarity
            if selected_graph_type == "Select a graph type":
                st.sidebar.warning("Please select a graph type.")
            elif selected_graph_type not in ["Correlation Heatmap"] and not columns_to_plot:
                st.sidebar.warning(f"Please select appropriate columns for {selected_graph_type}.")
            else:
                with st.spinner(f"Generating {selected_graph_type.lower()}..."):
                    img_buffer, graph_desc = generate_and_display_graph(
                        st.session_state['df'], 
                        selected_graph_type.lower().replace(" ", "_"), 
                        columns_to_plot
                    )
                    if img_buffer:
                        # Add graph to messages for persistent display
                        st.session_state.messages.append({"role": "graph", "content": img_buffer, "caption": graph_desc})
                        
                        # Get AI interpretation of the graph
                        interpretation_prompt = (
                            f"A {selected_graph_type.lower()} of the dataset was just generated. "
                            f"Here is its description: '{graph_desc}'. "
                            "Please provide a concise interpretation of what this graph tells us about the data, "
                            "especially in the context of data preprocessing. Do NOT provide code."
                        )
                        append_debug_log(f"DEBUG: Graph interpretation prompt:\n{interpretation_prompt}\n---") # Debug print
                        ai_interpretation = generate_openai_response(interpretation_prompt)
                        append_debug_log(f"DEBUG: Graph interpretation response:\n{ai_interpretation}\n---") # Debug print
                        st.session_state.messages.append({"role": "assistant", "content": ai_interpretation})
                        
                        # Add graph and its interpretation to report content
                        st.session_state.report_content.append({"type": "heading", "level": 2, "content": f"Graph Generated: {selected_graph_type}"})
                        st.session_state.report_content.append({"type": "image", "content": img_buffer, "caption": graph_desc})
                        st.session_state.report_content.append({"type": "text", "content": ai_interpretation}) # Add interpretation to report
                    else:
                        st.session_state.messages.append({"role": "assistant", "content": graph_desc}) # graph_desc will contain error message
                        st.session_state.report_content.append({"type": "text", "content": graph_desc})
                st.rerun() # Rerun to display the new message/graph immediately

        # --- Statistical Tests (New Section) ---
        st.sidebar.markdown("---")
        st.sidebar.header("Perform Statistical Tests")

        # UPDATED: Added all new statistical test options
        test_options = [
            "Select a test", 
            "Cronbach’s Alpha (Reliability)", # NEW
            "Z-Test (Two Sample Means)", # NEW
            "ANOVA", 
            "Independent T-test", # Assumes Equal Variances
            "Paired T-test", 
            "Chi-squared Test", 
            "Pearson Correlation(Validity-Linear)", 
            "Spearman Rank Correlation(Validity – Monotonic)",
            "F-Test Two-Sample for Variances", 
            "T-test: Two-Sample Assuming Unequal Variances", 
            "Shapiro-Wilk Test (Normality)" 
        ]
        selected_test = st.sidebar.selectbox("Choose Statistical Test:", test_options, key="stat_test_select")

        stat_col1 = None
        stat_col2 = None
        stat_cols_multiselect = None # For Cronbach's Alpha

        # Dynamic column selection based on selected test
        if selected_test == "Cronbach’s Alpha (Reliability)": # NEW
            st.sidebar.info("Cronbach’s Alpha: Measures internal consistency of multiple numerical items (e.g., survey questions).")
            stat_cols_multiselect = st.sidebar.multiselect("Select Numerical Items (2+ columns):", numerical_columns, key="cronbach_cols")
        elif selected_test == "Z-Test (Two Sample Means)": # NEW
            st.sidebar.info("Z-Test: Compares means of a numerical variable between 2 independent groups (large sample or known population SD).")
            stat_col1 = st.sidebar.selectbox("Numerical Variable:", ["Select column"] + numerical_columns, key="ztest_num_col")
            stat_col2 = st.sidebar.selectbox("Grouping Variable (2 categories):", ["Select column"] + categorical_columns, key="ztest_cat_col")
        elif selected_test == "ANOVA":
            st.sidebar.info("ANOVA: Compares means of a numerical variable across 2+ categories.")
            stat_col1 = st.sidebar.selectbox("Numerical Variable (Dependent):", ["Select column"] + numerical_columns, key="anova_num_col")
            stat_col2 = st.sidebar.selectbox("Categorical Variable (Independent):", ["Select column"] + categorical_columns, key="anova_cat_col")
        elif selected_test == "Independent T-test":
            st.sidebar.info("T-test (Equal Variances): Compares means of a numerical variable between 2 groups.")
            stat_col1 = st.sidebar.selectbox("Numerical Variable:", ["Select column"] + numerical_columns, key="ttest_eq_num_col")
            stat_col2 = st.sidebar.selectbox("Grouping Variable (2 categories):", ["Select column"] + categorical_columns, key="ttest_eq_cat_col")
        elif selected_test == "Paired T-test":
            st.sidebar.info("Paired T-test: Compares means of two related numerical variables (e.g., before/after).")
            stat_col1 = st.sidebar.selectbox("Numerical Variable 1 (e.g., Before):", ["Select column"] + numerical_columns, key="paired_ttest_num1_col")
            stat_col2 = st.sidebar.selectbox("Numerical Variable 2 (e.g., After):", ["Select column"] + numerical_columns, key="paired_ttest_num2_col")
        elif selected_test == "Chi-squared Test":
            st.sidebar.info("Chi-squared: Tests association between two categorical variables.")
            stat_col1 = st.sidebar.selectbox("Categorical Variable 1:", ["Select column"] + categorical_columns, key="chi2_cat1_col")
            stat_col2 = st.sidebar.selectbox("Categorical Variable 2:", ["Select column"] + categorical_columns, key="chi2_cat2_col")
        elif selected_test == "Pearson Correlation":
            st.sidebar.info("Pearson Correlation: Measures linear relationship between two numerical variables.")
            stat_col1 = st.sidebar.selectbox("Numerical Variable 1:", ["Select column"] + numerical_columns, key="pearson_num1_col")
            stat_col2 = st.sidebar.selectbox("Numerical Variable 2:", ["Select column"] + numerical_columns, key="pearson_num2_col")
        elif selected_test == "Spearman Rank Correlation":
            st.sidebar.info("Spearman Rank Correlation: Measures monotonic relationship between two numerical/ordinal variables.")
            stat_col1 = st.sidebar.selectbox("Numerical/Ordinal Variable 1:", ["Select column"] + numerical_columns, key="spearman_num1_col")
            stat_col2 = st.sidebar.selectbox("Numerical/Ordinal Variable 2:", ["Select column"] + numerical_columns, key="spearman_num2_col")
        elif selected_test == "F-Test Two-Sample for Variances":
            st.sidebar.info("F-Test: Compares the variances of two independent numerical samples.")
            stat_col1 = st.sidebar.selectbox("Numerical Sample 1:", ["Select column"] + numerical_columns, key="f_test_num1_col")
            stat_col2 = st.sidebar.selectbox("Numerical Sample 2:", ["Select column"] + numerical_columns, key="f_test_num2_col")
        elif selected_test == "T-test: Two-Sample Assuming Unequal Variances":
            st.sidebar.info("T-test (Unequal Variances): Compares means of a numerical variable between 2 groups, not assuming equal variances.")
            stat_col1 = st.sidebar.selectbox("Numerical Variable:", ["Select column"] + numerical_columns, key="ttest_uneq_num_col")
            stat_col2 = st.sidebar.selectbox("Grouping Variable (2 categories):", ["Select column"] + categorical_columns, key="ttest_uneq_cat_col")
        elif selected_test == "Shapiro-Wilk Test (Normality)":
            st.sidebar.info("Shapiro-Wilk Test: Tests if a numerical variable is normally distributed.")
            stat_col1 = st.sidebar.selectbox("Numerical Variable:", ["Select column"] + numerical_columns, key="shapiro_num_col") # Renamed from stat_col_single to stat_col1 for consistency
        
        if st.sidebar.button(f"Run {selected_test}"):
            append_debug_log(f"DEBUG: Button '{selected_test}' clicked.")
            append_debug_log(f"DEBUG: selected_test='{selected_test}', stat_col1='{stat_col1}', stat_col2='{stat_col2}', stat_cols_multiselect='{stat_cols_multiselect}'")

            # Validate column selections based on test type
            is_valid_selection = True
            if selected_test == "Select a test":
                st.sidebar.warning("Please select a statistical test to run.")
                is_valid_selection = False
            elif selected_test == "Cronbach’s Alpha (Reliability)": # NEW validation for multi-select
                if not stat_cols_multiselect or len(stat_cols_multiselect) < 2:
                    st.sidebar.warning("Please select at least two numerical columns for Cronbach’s Alpha.")
                    is_valid_selection = False
                else: # Pass the list of columns as col1
                    stat_col1 = stat_cols_multiselect
                    stat_col2 = None # Ensure col2 is None for this test
            elif selected_test == "Shapiro-Wilk Test (Normality)":
                if stat_col1 == "Select column":
                    st.sidebar.warning("Please select a numerical variable for the Shapiro-Wilk Test.")
                    is_valid_selection = False
            elif selected_test in ["ANOVA", "Independent T-test", "Paired T-test", "Chi-squared Test",
                                   "Pearson Correlation", "Spearman Rank Correlation",
                                   "F-Test Two-Sample for Variances", "T-test: Two-Sample Assuming Unequal Variances", "Z-Test (Two Sample Means)"]: # Added Z-Test here
                if stat_col1 == "Select column" or stat_col2 == "Select column":
                    st.sidebar.warning("Please select all required columns for the chosen test.")
                    is_valid_selection = False
            
            if is_valid_selection:
                append_debug_log(f"DEBUG: Calling perform_statistical_test for {selected_test} with col1={stat_col1}, col2={stat_col2}")
                with st.spinner(f"Running {selected_test}..."):
                    # Convert selected_test to the internal string used in perform_statistical_test
                    internal_test_type = selected_test.replace(" ", "_").replace("(", "").replace(")", "").lower().replace("’", "")
                    
                    test_results_str, structured_results_for_ui, test_error, plot_buffer_stat, plot_caption_stat = perform_statistical_test(
                        st.session_state['df'], 
                        internal_test_type, 
                        stat_col1, # This will be list of columns for Cronbach's, single for others
                        stat_col2
                    )

                    if test_error:
                        st.session_state.messages.append({"role": "assistant", "content": test_error})
                        st.session_state.report_content.append({"type": "text", "content": f"Statistical Test Error ({selected_test}): {test_error}"})
                    else:
                        # Add the initial text summary of the test to messages
                        st.session_state.messages.append({"role": "assistant", "content": test_results_str})
                        st.session_state.report_content.append({"type": "heading", "level": 2, "content": f"Statistical Test: {selected_test}"})
                        st.session_state.report_content.append({"type": "text", "content": test_results_str})

                        # Display structured results in UI and add to report based on test type
                        if selected_test in ["Independent T-test", "Paired T-test", "Z-Test (Two Sample Means)", "T-test: Two-Sample Assuming Unequal Variances"] and structured_results_for_ui is not None:
                            group_stats_df, test_stats_df = structured_results_for_ui # Unpack the tuple
                            
                            st.session_state.messages.append({"role": "dataframe", "title": "Group Statistics", "content": group_stats_df})
                            st.session_state.report_content.append({"type": "stat_table", "title": "Group Statistics", "dataframe": group_stats_df})

                            st.session_state.messages.append({"role": "dataframe", "title": "Test Results", "content": test_stats_df})
                            st.session_state.report_content.append({"type": "stat_table", "title": "Test Results", "dataframe": test_stats_df})
                        
                        elif selected_test == "ANOVA" and structured_results_for_ui is not None:
                            anova_df = structured_results_for_ui
                            st.session_state.messages.append({"role": "dataframe", "title": "ANOVA Summary Table", "content": anova_df})
                            st.session_state.report_content.append({"type": "stat_table", "title": "ANOVA Summary Table", "dataframe": anova_df})

                        elif selected_test in ["Pearson Correlation", "Spearman Rank Correlation", "F-Test Two-Sample for Variances", 
                                               "Shapiro-Wilk Test (Normality)", "Cronbach’s Alpha (Reliability)"] and structured_results_for_ui is not None: # UPDATED for Cronbach's
                            # These tests return a single DataFrame
                            single_table_df = structured_results_for_ui
                            table_title = f"{selected_test} Results"
                            if selected_test == "Shapiro-Wilk Test (Normality)":
                                table_title = f"Shapiro-Wilk Test Results for {stat_col1}"
                            
                            st.session_state.messages.append({"role": "dataframe", "title": table_title, "content": single_table_df})
                            st.session_state.report_content.append({"type": "stat_table", "title": table_title, "dataframe": single_table_df})

                        elif selected_test == "Chi-squared Test" and structured_results_for_ui is not None:
                            observed_df, expected_df, chi2_val, p_val, dof_val = structured_results_for_ui

                            st.session_state.messages.append({"role": "dataframe", "title": "Observed Frequencies", "content": observed_df})
                            st.session_state.report_content.append({"type": "stat_table", "title": "Observed Frequencies", "dataframe": observed_df})

                            st.session_state.messages.append({"role": "dataframe", "title": "Expected Frequencies", "content": expected_df})
                            st.session_state.report_content.append({"type": "stat_table", "title": "Expected Frequencies", "dataframe": expected_df})

                            # Add summary text for Chi-squared after tables
                            chi2_summary_text = (
                                f"Chi-squared Test Statistics:\n"
                                f"  Chi-squared statistic: {chi2_val:.4f}\n"
                                f"  P-value: {p_val:.4f}\n"
                                f"  Degrees of Freedom (dof): {dof_val}\n"
                                "Interpretation will be provided by the AI."
                            )
                            st.session_state.messages.append({"role": "assistant", "content": chi2_summary_text})
                            st.session_state.report_content.append({"type": "text", "content": chi2_summary_text})

                        # Add plot if the test generated one (currently only for regression, but structure is here)
                        if plot_buffer_stat and plot_caption_stat:
                            st.session_state.messages.append({"role": "graph", "content": plot_buffer_stat, "caption": plot_caption_stat})
                            st.session_state.report_content.append({"type": "image", "content": plot_buffer_stat, "caption": plot_caption_stat})


                        # Get AI interpretation of the test results
                        interpretation_prompt = (
                            f"A {selected_test} was just performed with the following results:\n"
                            f"{test_results_str}\n" # Use the original results_str for AI prompt
                            "Please provide a concise, plain-language interpretation of these results, "
                            "focusing on what the p-value means and the implications for the relationship between the variables. "
                            "Do NOT provide code or markdown formatting."
                        )
                        append_debug_log(f"DEBUG: Stat test interpretation prompt:\n{interpretation_prompt}\n---") # Debug print
                        ai_interpretation = generate_openai_response(interpretation_prompt)
                        append_debug_log(f"DEBUG: Stat test interpretation response:\n{ai_interpretation}\n---") # Debug print
                        st.session_state.messages.append({"role": "assistant", "content": ai_interpretation})
                        st.session_state.report_content.append({"type": "text", "content": ai_interpretation})
                st.rerun()

        # --- NEW: Regression Analysis Section ---
        st.sidebar.markdown("---")
        st.sidebar.header("Regression Analysis")

        regression_dep_var = st.sidebar.selectbox("Dependent Variable (Y):", ["Select column"] + numerical_columns, key="reg_dep_var")
        regression_indep_var = st.sidebar.selectbox("Independent Variable (X):", ["Select column"] + numerical_columns, key="reg_indep_var")

        if st.sidebar.button("Run Simple Linear Regression"):
            if regression_dep_var == "Select column" or regression_indep_var == "Select column":
                st.sidebar.warning("Please select both dependent and independent variables for regression.")
            else:
                with st.spinner("Running Simple Linear Regression..."):
                    reg_summary, reg_plot_buffer, reg_plot_caption, reg_error = perform_linear_regression(
                        st.session_state['df'], 
                        regression_dep_var, 
                        regression_indep_var
                    )
                    
                    if reg_error:
                        st.session_state.messages.append({"role": "assistant", "content": reg_error})
                        st.session_state.report_content.append({"type": "text", "content": f"Regression Error: {reg_error}"})
                    else:
                        st.session_state['regression_results'] = reg_summary # Store for display and report
                        st.session_state['regression_plot_buffer'] = reg_plot_buffer
                        st.session_state['regression_plot_caption'] = reg_plot_caption

                        # Display regression plot in chat
                        st.session_state.messages.append({"role": "graph", "content": reg_plot_buffer, "caption": reg_plot_caption})
                        st.session_state.report_content.append({"type": "heading", "level": 2, "content": "Simple Linear Regression"})
                        st.session_state.report_content.append({"type": "image", "content": reg_plot_buffer, "caption": reg_plot_caption})

                        # Display regression summary in chat (as code block for formatting)
                        st.session_state.messages.append({"role": "code", "title": "Regression Model Summary", "content": reg_summary})
                        st.session_state.report_content.append({"type": "regression_summary", "content": reg_summary})

                        # Get AI interpretation of regression results
                        interpretation_prompt = (
                            f"A simple linear regression was performed with dependent variable '{regression_dep_var}' and independent variable '{regression_indep_var}'. "
                            f"Here is the model summary:\n{reg_summary}\n\n"
                            "Please provide a concise, plain-language interpretation of these regression results, "
                            "focusing on the relationship between the variables, the significance of the model, and the R-squared value. "
                            "Do NOT provide code or markdown formatting."
                        )
                        append_debug_log(f"DEBUG: Regression interpretation prompt:\n{interpretation_prompt}\n---")
                        ai_interpretation = generate_openai_response(interpretation_prompt)
                        append_debug_log(f"DEBUG: Regression interpretation response:\n{ai_interpretation}\n---")
                        st.session_state.messages.append({"role": "assistant", "content": ai_interpretation})
                        st.session_state.report_content.append({"type": "text", "content": ai_interpretation})
                st.rerun()

        # Chat input box - positioned before the footer
        if prompt := st.chat_input("Ask about preprocessing or analysis..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            # Update user goal if explicitly stated in the prompt
            prompt_lower = prompt.lower()
            if "my goal is" in prompt_lower or "i want to do" in prompt_lower or "my objective is" in prompt_lower:
                st.session_state['user_goal'] = prompt # Simple capture for now
                st.session_state.report_content.append({"type": "heading", "level": 2, "content": "User's Stated Goal"})
                st.session_state.report_content.append({"type": "text", "content": prompt})

            # Construct prompt for OpenAI, including data summary and full chat history
            full_prompt = (
                f"Dataset Summary:\n{st.session_state['data_summary_text']}\n\n"
                "Conversation History:\n" + "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages if m["role"] in ["user", "assistant"]]) + # Filter for AI
                f"\n\nUser's current message: {prompt}\n\n"
                "Based on the dataset summary and our conversation, provide tailored preprocessing advice, "
                "including explanations. Do NOT provide Python code snippets. "
                "If the user has stated a goal, ensure your advice aligns with it."
            )
            append_debug_log(f"DEBUG: General chat prompt:\n{full_prompt}\n---") # Debug print
            with st.spinner("Generating response..."):
                response = generate_openai_response(full_prompt)
                st.session_state.messages.append({"role": "assistant", "content": response})
                st.session_state.report_content.append({"type": "text", "content": response}) # Log AI responses for report
            st.rerun() # Rerun to display the new message/graph

    st.sidebar.markdown("---")
    st.sidebar.header("Report & Actions")

    # Add a "Reset Chat" button in the sidebar
    if st.sidebar.button("Reset Chat"):
        st.session_state['messages'] = []
        st.session_state['report_content'] = []
        st.session_state['user_goal'] = "Not specified"
        st.session_state['debug_logs'] = [] # Clear debug logs on chat reset
        st.session_state['regression_results'] = None # Reset regression results
        st.session_state['regression_plot_buffer'] = None # Reset regression plot
        st.session_state['regression_plot_caption'] = None # Reset regression plot caption

        # If a file is uploaded, re-trigger initial analysis
        if st.session_state['df'] is not None:
            summary_text, summary_table = get_data_summary(st.session_state['df'])
            st.session_state['data_summary_text'] = summary_text
            st.session_state['data_summary_table'] = summary_table
            st.session_state['report_content'].append({"type": "heading", "level": 2, "content": "Dataset Overview"})
            overview_text_end_index = summary_text.find("Column Details:")
            if overview_text_end_index != -1:
                overview_text = summary_text[:overview_text_end_index].strip()
            else:
                overview_text = st.session_state['data_summary_text'].split("Column Details:")[0].strip() # Fallback
            st.session_state['report_content'].append({"type": "text", "content": overview_text})
            st.session_state['report_content'].append({"type": "table", "headers": summary_table[0], "rows": summary_table[1:]})
            
            initial_ai_prompt = (
                "Here is a detailed summary of the user's dataset:\n\n"
                f"{st.session_state['data_summary_text']}\n\n"
                "Based on this, what are the initial preprocessing considerations? "
                "Please also ask the user about their primary goal (e.g., classification, regression, exploratory analysis) for this dataset."
            )
            append_debug_log(f"DEBUG: Reset chat initial AI prompt:\n{initial_ai_prompt}\n---") # Debug print
            with st.spinner("Analyzing data and generating initial insights..."):
                initial_response = generate_openai_response(initial_ai_prompt)
                st.session_state.messages.append({"role": "assistant", "content": initial_response})
                st.session_state.report_content.append({"type": "heading", "level": 2, "content": "Initial Preprocessing Considerations"})
                st.session_state.report_content.append({"type": "text", "content": initial_response})
        st.rerun()


    # Download Report Button
    if st.sidebar.button("Generate & Download Report"):
        if st.session_state['df'] is not None:
            with st.spinner("Generating report..."):
                report_buffer = create_report_doc(st.session_state['report_content'])
                st.sidebar.download_button(
                    label="Download Report (.docx)",
                    data=report_buffer,
                    file_name="Data_Preprocessing_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_report_button"
                )
        else:
            st.sidebar.warning("Please upload a dataset first to generate a report.")

    # --- Footer ---
    st.markdown("---")
    # Centering the footer
    st.markdown(
        "<div style='text-align: center;'>"
        "SSO Consultants © 2025 | All Rights Reserved."
        "</div>",
        unsafe_allow_html=True
    )

    # --- In-App Debug Logs ---
    st.expander_debug = st.expander("Show Debug Logs")
    with st.expander_debug:
        if st.button("Clear Debug Logs", key="clear_debug_logs_button"): # Removed backslash
            st.session_state['debug_logs'] = []
            st.rerun()
        for log_entry in st.session_state['debug_logs']:
            st.code(log_entry, language='text')

# --- Run the App ---
if not st.session_state['logged_in']:
    check_password()
else:
    main_app()
